"""Extract plain text from resume uploads (PDF / DOCX / text)."""

from __future__ import annotations

import io
import re


class ExtractError(ValueError):
    pass


# Multi-word section titles — case-insensitive. Safe mid-stream (rarely body prose).
_SECTION_MULTI_RE = re.compile(
    r"(?i)(?<![A-Za-z])("
    r"professional\s+summary|career\s+summary|about\s+me|"
    r"technical\s+skills|core\s+competencies|tech\s+stack|"
    r"professional\s+experience|work\s+experience|employment\s+history|work\s+history|"
    r"selected\s+projects|personal\s+projects|key\s+projects|"
    r"academic\s+background|additional\s+information|"
    r"volunteer(?:ing)?\s+experience|open\s+source(?:\s+contributions)?"
    r")(?![A-Za-z])"
)

# Single-token ALL-CAPS headings only (typical PDF chrome). Never match mid-sentence
# lowercase "experience". Multi-word phrases are placeholder-protected first so
# SUMMARY inside PROFESSIONAL SUMMARY is not re-split.
_SECTION_ALLCAPS_RE = re.compile(
    r"(?<![A-Za-z])("
    r"SUMMARY|PROFILE|OBJECTIVE|SKILLS|TECHNOLOGIES|"
    r"EXPERIENCE|EMPLOYMENT|"
    r"PROJECTS|PORTFOLIO|"
    r"EDUCATION|ACADEMICS|"
    r"CERTIFICATIONS|CERTIFICATES|LICENSES|AWARDS|HONORS|ACHIEVEMENTS|PUBLICATIONS|"
    r"INTERESTS|HOBBIES|ACTIVITIES|LEADERSHIP|RESEARCH|"
    r"REFERENCES"
    r")(?![A-Za-z])"
)

# Title-Case single-token headings only after end-of-sentence punctuation.
_SECTION_TITLECASE_RE = re.compile(
    r"(?<=[.!?])\s+("
    r"Summary|Profile|Objective|"
    r"Skills|Technologies|"
    r"Experience|Employment|"
    r"Projects|Portfolio|"
    r"Education|"
    r"Certifications|Awards|"
    r"References"
    r")(?=\s+[A-Z0-9(])"
)

# Role title + date range stuck mid-stream after section flatten
_TITLE_DATE_RE = re.compile(
    r"(?i)(?<=\S)\s+("
    r"(?:junior|senior|lead|staff|principal|full[- ]?stack|frontend|front-end|"
    r"backend|back-end|software|data|ml|ai|mobile|devops|cloud|security|product|"
    r"web|platform|systems?|intern)"
    r"[A-Za-z0-9 /|&+.,'-]{0,60}?"
    r"\s+(?:\d{1,2}[/.-]\d{4}|\d{4}|"
    r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4})"
    r"\s*(?:[-–—]|to)\s*"
    r"(?:present|current|now|\d{1,2}[/.-]\d{4}|\d{4}|"
    r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4})"
    r")"
)

_KNOWN_LINE_HEADING = re.compile(
    r"(?i)^(professional\s+summary|career\s+summary|summary|profile|objective|"
    r"technical\s+skills|core\s+competencies|skills|technologies|"
    r"professional\s+experience|work\s+experience|employment\s+history|"
    r"experience|employment|selected\s+projects|projects|"
    r"education|certifications|references)\s*:?\s*$"
)

# Skill category rows that appear as "Languages: foo Frontend: bar" when lines were lost
_SKILL_CATEGORY_IN_STREAM_RE = re.compile(
    r"(?i)(?<=\S)\s+("
    r"Languages?|Frontend|Backend(?:\s+and\s+APIs?)?|Databases?(?:\s+and\s+Vector\s+Search)?|"
    r"DevOps(?:\s+and\s+Practices)?|Tools|Frameworks|Cloud|Mobile|Testing|Libraries"
    r")\s*:\s*"
)


def extract_text(data: bytes, filename: str) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _pdf(data)
    if name.endswith(".docx"):
        return _docx(data)
    try:
        text = data.decode("utf-8", errors="ignore")
    except Exception as exc:  # noqa: BLE001
        raise ExtractError("Unsupported or unreadable text file") from exc
    return _normalize(text)


def needs_structure_restore(text: str) -> bool:
    """True when text looks like a flattened resume (lost section line breaks)."""
    if not text or not text.strip():
        return False
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if len(lines) <= 3:
        return True
    heading_lines = sum(1 for ln in lines if _KNOWN_LINE_HEADING.match(ln))
    if heading_lines < 2 and len(text) > 400:
        return True
    return False


def restore_resume_structure(text: str) -> str:
    """
    Heal resumes whose section breaks were flattened to spaces.

    Some PDF extractors (or rewrites) drop newlines entirely. The structured
    pipeline expects one heading / entry per line, so we re-insert breaks before
    known section titles and role title+date patterns — without treating bare
    body words like "experience" as headings.
    """
    if not text or not text.strip():
        return text or ""
    if not needs_structure_restore(text):
        return text

    out = text
    protected: list[str] = []

    def _protect(match: re.Match[str]) -> str:
        protected.append(_title_heading(match.group(1)))
        return f"\n@@HEAD{len(protected) - 1}@@\n"

    out = _SECTION_MULTI_RE.sub(_protect, out)
    out = _SECTION_ALLCAPS_RE.sub(lambda m: "\n" + m.group(1).upper() + "\n", out)
    out = _SECTION_TITLECASE_RE.sub(lambda m: "\n" + m.group(1) + "\n", out)

    for i, heading in enumerate(protected):
        out = out.replace(f"@@HEAD{i}@@", heading)

    out = _TITLE_DATE_RE.sub(r"\n\1", out)
    out = _SKILL_CATEGORY_IN_STREAM_RE.sub(r"\n\1: ", out)
    out = re.sub(r"\s+\|\s+", " | ", out)
    out = re.sub(r"[ \t]+", " ", out)
    out = re.sub(r" *\n *", "\n", out)
    out = re.sub(r"\n{3,}", "\n\n", out)
    return out.strip()


def _title_heading(raw: str) -> str:
    return re.sub(r"\s+", " ", raw.strip()).upper()


def _normalize(text: str) -> str:
    text = text.replace("\x00", " ")
    # Preserve newlines; only collapse horizontal whitespace per line
    text = "\n".join(re.sub(r"[ \t]+", " ", ln).strip() for ln in text.splitlines())
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = restore_resume_structure(text)
    text = text.strip()
    if len(text) < 40:
        raise ExtractError("Resume text too short or unreadable")
    return text


def _pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise ExtractError("PDF support unavailable") from exc
    try:
        reader = PdfReader(io.BytesIO(data))
        parts = [(p.extract_text() or "") for p in reader.pages]
    except Exception as exc:  # noqa: BLE001
        raise ExtractError("Could not read this PDF") from exc
    return _normalize("\n".join(parts))


def _docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise ExtractError("DOCX support unavailable — install python-docx") from exc
    try:
        doc = Document(io.BytesIO(data))
        parts: list[str] = []
        for p in doc.paragraphs:
            t = (p.text or "").strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
    except Exception as exc:  # noqa: BLE001
        raise ExtractError("Could not read this DOCX") from exc
    return _normalize("\n".join(parts))
